#!/usr/bin/env python
# -*- coding: utf-8 -*-

#py -m pip install python-docx
import docx
from docx.shared import Inches
import sys
import os
import re
import json
import hashlib
import html

'''
from wordinserter import insert, parse
from comtypes.client import CreateObject
#py -m pip install wordinserter
#py -m pip install comtypes
# This opens Microsoft Word and creates a new document.
word = CreateObject("Word.Application")
word.Visible = True # Don't set this to True in production!
document = word.Documents.Add()
'''

from comtypes.gen import Word as constants

if sys.version_info[0] == 2:
    from urllib2 import urlopen  # Python 2
else:
    from urllib.request import urlopen  # Python3
    import urllib.request
    import urllib.parse

def http_get(url):
    req = urllib.request.Request(url)
    response = urllib.request.urlopen(req,timeout = 60)
    data = response.read().decode()
    print("http_get",url,len(data.encode('utf-8')))
    return data

def get_url(url):
    data = urlopen(url).read()
    return data.decode("utf-8")

def save_file(videopath,url):
    data = urlopen(url).read()
    #print("get data finish ",videopath,url)
    vfile = open(videopath, 'wb')
    vfile.write(data)
    vfile.close()

def find_link_html(htmll):
    pattern = re.compile( r'<a [^<>]*href="([^"]*)"') 
    result = pattern.findall(htmll)

    return result

def find_media_url(htmll):
    pattern = re.compile( r'<img [^<>]*src="([^"]*)"') 
    result1 = pattern.findall(htmll)
    pattern = re.compile( r'<source [^<>]*src="([^"]*)"') 
    result2 = pattern.findall(htmll)

    result = result1 + result2

    return result

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

username = 'news'
url = "http://f.hayoou.com/timline/search.php?username="+username+"&s=&page=1&size=100&json=1"
#http://f.hayoou.com/timline/search.php?username=news&s=&page=1&size=10&all=1
jsondata = http_get(url)
#jsonobj=json.dumps(jsondata, ensure_ascii=False).encode('utf-8')
#jsondata = jsondata[jsondata.find("{"):]
#print(jsondata.encode('utf-8'))
jsonObj = json.loads(jsondata)
#新建文档
doc_new = docx.Document()
os.makedirs("./doc/"+username+"/files/artical", mode=0o777, exist_ok=True)

#print(jsonObj)
count = 0
for timelineobj in jsonObj:
    htmll = timelineobj['content']

    count+=1
    print("---------------")
    print(count,"%.2f"%(count/len(jsonObj)*100),"%")

    doc_new.add_paragraph(timelineobj['date'], style='ListBullet')

    '''
    operations = parse(htmll, parser="html")
    insert(operations, document=document, constants=constants)
    '''

    text =htmll
    textlist= htmll.split("<span ")
    if len(textlist)>0:
        text = textlist[0]

    textlist= text.split("<a ")
    if len(textlist)>0:
        text = textlist[0]

    rc = re.compile("<.*?>" )
    text = rc.sub('',text)
    p = doc_new.add_paragraph( text )
    print(text)

    has_artical = (htmll.find("<table") != -1)

    media_url =find_media_url(htmll)
    counter = 0
    if len(media_url)>0:
        path0 = "./doc/"+username+"/files/artical/"
        if not has_artical:
            os.makedirs("./doc/"+username+"/files/"+timelineobj['date'][0:10], mode=0o777, exist_ok=True)
            path0 = "./doc/"+username+"/files/"+timelineobj['date'][0:10]+"/"

        print("len(media_url)",len(media_url))
        for url in media_url:
            if url=="":
                continue
            
            url1 = html.unescape(url)
            url1=url1.replace("&req=imageView2/1/w/240/h/240","")
            url1a = url1.split("?imageMogr2")
            if len(url1a)>0:
                url1 =url1a[0]
            #print("url",url,"url1",url1)

            counter +=1
            if url1.find("video.qq.com")==-1:
                if has_artical:
                    mediafile = timelineobj['date'] +" "+ str(counter) + ".jpg"
                else:
                    mediafile = timelineobj['date'] +" "+ str(counter) + ".png"
            else:
                mediafile = timelineobj['date'] +" "+ str(counter) + ".mp4"

            mediafile = mediafile.replace(":","-")
            mediafilepath = path0+mediafile
            download_ok = 0

            if os.path.exists(mediafilepath): 
                #print("exists",mediafilepath)
                download_ok=1
            elif os.path.exists(mediafilepath+".error"):
                pass
            else:
                try:
                    save_file(mediafilepath,url1)
                    download_ok=1
                except Exception as e:
                    save_file(mediafilepath+".error","https://boxmy.hayoou.com/1563938514578/empty2.png")

            if download_ok==1:
                try:
                    p = doc_new.add_picture(mediafilepath, width=Inches(1.25))
                    if( not has_artical ):
                        hyperlink = add_hyperlink(p, url1, '链接',   'FF8822', False)

                except Exception as e:
                    print("add_picture fail: ",mediafilepath)
            else:
                print("not download ok",url1)
    else:
        p = doc_new.add_paragraph(  )


    links =  find_link_html(htmll)
    if len(links)>0 and has_artical:
        #p = doc_new.add_paragraph( )
        text2 = htmll
        start = text2.find("<span ")
        if start == -1:
            start = 99999
        start1 = text2.find("<a ")
        if start1 == -1:
            start1 = 99999

        start  = min(start,start1)

        if start != 99999:
            text2 = text2[start:]

        rc = re.compile("<.*?>" )
        content = rc.sub('',text2)
        rc = re.compile("[\r\n]*" )
        content = rc.sub('',content)

        print("content",content,links[0])

        for link in links:
            if link=="":
                continue
            if len(content.replace(" ",""))>1:
                try:
                    hyperlink = add_hyperlink(p, link, content,   'FF8822', False)
                except Exception as e:
                    p = doc_new.add_paragraph(  )
                    hyperlink = add_hyperlink(p, link, content,   'FF8822', False)
            break


doc_new.save("./doc/"+username+"/"+username+".docx")

